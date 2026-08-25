# utils.py
from datetime import datetime, timedelta, date
from functools import wraps
from flask import session, flash, redirect, url_for
import os
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ==================== الصلاحيات ====================
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user_id' not in session:
            flash('يجب تسجيل الدخول أولاً', 'warning')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

def role_required(*roles):
    def decorator(f):
        @wraps(f)
        def decorated_function(*args, **kwargs):
            if session.get('role') not in roles:
                flash('غير مصرح لك بالدخول لهذه الصفحة', 'danger')
                return redirect(url_for('dashboard'))
            return f(*args, **kwargs)
        return decorated_function
    return decorator

# ==================== صلاحية التعديل ====================
def can_edit(user_role, record_date, record_created_by=None, current_user_id=None):
    if user_role in ['meg', 'admin', 'mariam']:
        return True
    if user_role == 'rehab':
        days_diff = (datetime.now().date() - record_date).days
        return days_diff <= 7
    if user_role in ['mohamed', 'ahmed', 'eid', 'abdo']:
        if record_created_by is not None and current_user_id is not None:
            if record_created_by != current_user_id:
                return False
        hours_diff = (datetime.now() - datetime.combine(record_date, datetime.min.time())).total_seconds() / 3600
        return hours_diff <= 24
    return False

# ==================== توليد ملفات Excel ====================
def get_excel_dir():
    base_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'excel_files')
    os.makedirs(base_dir, exist_ok=True)
    return base_dir

def create_or_load_workbook(file_name, headers):
    file_path = os.path.join(get_excel_dir(), file_name)
    if os.path.exists(file_path):
        wb = load_workbook(file_path)
        ws = wb.active
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = "التقرير"
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True, color='FFFFFF', size=12)
            cell.fill = PatternFill(start_color='7C3AED', end_color='7C3AED', fill_type='solid')
            cell.alignment = Alignment(horizontal='center', vertical='center')
            cell.border = Border(bottom=Side(style='thin'))
    return wb, ws, file_path

def add_row_to_excel(file_name, headers, row_data):
    wb, ws, file_path = create_or_load_workbook(file_name, headers)
    ws.append(row_data)
    wb.save(file_path)
    return file_path

# ==================== توليد تقارير Word ====================
def generate_word_report(company_name, report_title, report_date, period, table_headers, table_data, output_path):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    company_paragraph = doc.add_paragraph()
    company_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = company_paragraph.add_run(company_name)
    run.font.size = Pt(20)
    run.font.bold = True

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_paragraph.add_run(report_title)
    run.font.size = Pt(16)
    run.font.bold = True

    info_paragraph = doc.add_paragraph()
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_text = f"التاريخ: {report_date}    المدة: {period}"
    run = info_paragraph.add_run(info_text)
    run.font.size = Pt(12)

    doc.add_paragraph('─' * 50)

    if table_data:
        table = doc.add_table(rows=1, cols=len(table_headers))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(table_headers):
            hdr_cells[i].text = str(header)
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True

        for row in table_data:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = str(value)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run(f"© {company_name} - {datetime.now().year}").font.size = Pt(10)

    doc.save(output_path)
    return output_path

# ==================== الربط مع Google Drive ====================
class GoogleDriveService:
    def __init__(self, credentials_file='client_secrets.json', folder_id=None):
        self.credentials_file = credentials_file
        self.folder_id = folder_id
        self.service = None
        self._init_service()

    def _init_service(self):
        try:
            if not os.path.exists(self.credentials_file):
                print("ℹ️ Google Drive service disabled (client_secrets.json not found)")
                return

            from google.oauth2 import service_account
            from googleapiclient.discovery import build

            credentials = service_account.Credentials.from_service_account_file(
                self.credentials_file,
                scopes=['https://www.googleapis.com/auth/drive.file']
            )
            self.service = build('drive', 'v3', credentials=credentials)
        except Exception as e:
            print(f"Drive service error: {e}")
            self.service = None

    def upload_file(self, file_path, file_name):
        if not self.service or not self.folder_id:
            return None
        try:
            from googleapiclient.http import MediaFileUpload
            media = MediaFileUpload(file_path,
                                    mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            file_metadata = {
                'name': file_name,
                'parents': [self.folder_id]
            }
            file = self.service.files().create(body=file_metadata, media_body=media, fields='id').execute()
            return file.get('id')
        except Exception as e:
            print(f"Upload error: {e}")
            return None
